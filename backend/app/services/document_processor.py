"""Document processing service for text extraction."""

import io

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument

from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Service for extracting text from various document types."""

    SUPPORTED_MIME_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "application/vnd.ms-excel": "xls",
        "text/plain": "txt",
        "text/csv": "csv",
    }

    @classmethod
    def is_supported(cls, mime_type: str) -> bool:
        """Check if the MIME type is supported for text extraction."""
        return mime_type in cls.SUPPORTED_MIME_TYPES

    @classmethod
    def get_file_type(cls, mime_type: str) -> str | None:
        """Get the file type from MIME type."""
        return cls.SUPPORTED_MIME_TYPES.get(mime_type)

    async def extract_text(
        self,
        file_content: bytes,
        mime_type: str,
        filename: str,
    ) -> str:
        """
        Extract text from a document.

        Args:
            file_content: The raw file bytes
            mime_type: The MIME type of the file
            filename: The original filename

        Returns:
            Extracted text content

        Raises:
            ValueError: If the file type is not supported
        """
        logger.debug(
            "text_extraction_started",
            mime_type=mime_type,
            filename=filename,
        )

        file_type = self.get_file_type(mime_type)
        if not file_type:
            raise ValueError(f"Unsupported MIME type: {mime_type}")

        if file_type == "pdf":
            return self._extract_pdf(file_content)
        elif file_type == "docx":
            return self._extract_docx(file_content)
        elif file_type in ("xlsx", "xls"):
            return self._extract_excel(file_content)
        elif file_type in ("txt", "csv"):
            return self._extract_text_file(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using pdfplumber."""
        text_parts = []

        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Also extract table content
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            row_text = " | ".join(
                                str(cell) if cell else "" for cell in row
                            )
                            text_parts.append(row_text)

        return "\n\n".join(text_parts)

    def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        doc = DocxDocument(io.BytesIO(file_content))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    def _extract_excel(self, file_content: bytes) -> str:
        """Extract text from Excel files using pandas."""
        text_parts = []

        # Read all sheets
        excel_file = pd.ExcelFile(io.BytesIO(file_content))
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)

            # Add sheet name as header
            text_parts.append(f"=== Sheet: {sheet_name} ===")

            # Convert DataFrame to text
            for _, row in df.iterrows():
                row_values = [str(v) for v in row.values if pd.notna(v)]
                if row_values:
                    text_parts.append(" | ".join(row_values))

        return "\n\n".join(text_parts)

    def _extract_text_file(self, file_content: bytes) -> str:
        """Extract text from plain text or CSV files."""
        # Try different encodings
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                logger.debug(
                    "encoding_fallback",
                    encoding=encoding,
                )
                continue

        # Fallback with error handling
        logger.debug("encoding_using_fallback", encoding="utf-8_with_replace")
        return file_content.decode("utf-8", errors="replace")
